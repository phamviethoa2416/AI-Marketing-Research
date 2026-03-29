from __future__ import annotations

import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

# ─── Color constants (RGB tuples) ─────────────────────────────────────────────

NAVY   = RGBColor(0x1B, 0x2B, 0x4B)
BLUE   = RGBColor(0x25, 0x63, 0xEB)
ACCENT = RGBColor(0x0E, 0xA5, 0xE9)
GRAY   = RGBColor(0x64, 0x74, 0x8B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)


# ─── DOCX helpers ────────────────────────────────────────────────────────────

def set_run_color(run, color: RGBColor):
    run.font.color.rgb = color


def set_cell_background(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "E2E8F0")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_break(doc: Document):
    doc.add_page_break()


# ─── Style setup ──────────────────────────────────────────────────────────────

def configure_styles(doc: Document):
    styles = doc.styles

    # Normal
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    # Heading 1
    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    # Heading 2
    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = BLUE
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = NAVY
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)


# ─── Markdown parser ──────────────────────────────────────────────────────────

def parse_markdown_to_docx(doc: Document, markdown: str):
    lines = markdown.split("\n")
    i = 0
    bullet_batch: list[str] = []

    def flush_bullets():
        for b in bullet_batch:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(_strip_inline_md(b))
        bullet_batch.clear()



    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("### "):
            flush_bullets()
            _add_heading_para(doc, stripped[4:], level=3)

        elif stripped.startswith("## "):
            flush_bullets()
            add_horizontal_rule(doc)
            _add_heading_para(doc, stripped[3:], level=2)

        elif stripped.startswith("# "):
            flush_bullets()
            _add_heading_para(doc, stripped[2:], level=1)

        elif stripped.startswith(("- ", "* ")):
            bullet_batch.append(_strip_inline_md(stripped[2:]))

        elif stripped in ("---", "***"):
            flush_bullets()
            add_horizontal_rule(doc)

        elif not stripped:
            flush_bullets()

        else:
            flush_bullets()
            _add_inline_para(doc, stripped)

        i += 1

    flush_bullets()


def _add_heading_para(doc: Document, text: str, level: int):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    p = doc.add_paragraph(style=style_map.get(level, "Heading 1"))
    p.add_run(_strip_inline_md(text))


def _strip_inline_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _add_inline_para(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Split on bold/italic markers
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = p.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = p.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`") and token.endswith("`"):
            run = p.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            p.add_run(token)


# ─── Cover page ───────────────────────────────────────────────────────────────

def build_cover_page(doc: Document, pipeline_result: dict):
    title   = pipeline_result.get("report_title", "Báo cáo Nghiên cứu")
    query   = pipeline_result.get("query", "")
    quality = pipeline_result.get("quality", {})
    rid     = pipeline_result.get("report_id", "")

    # Logo / header block
    header_p = doc.add_paragraph()
    header_r = header_p.add_run("MULTI-AGENT RESEARCH SYSTEM  ·  BÁO CÁO NGHIÊN CỨU")
    header_r.font.name = "Calibri"
    header_r.font.size = Pt(10)
    header_r.font.color.rgb = ACCENT
    header_r.font.bold = True
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Title
    title_p = doc.add_paragraph()
    title_r = title_p.add_run(title.upper())
    title_r.font.name = "Calibri"
    title_r.font.size = Pt(24)
    title_r.font.bold = True
    title_r.font.color.rgb = NAVY
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Subtitle / query — bilingual
    if query:
        sub_p = doc.add_paragraph()
        sub_r = sub_p.add_run(f"Research query / Câu hỏi nghiên cứu:\n{query}")
        sub_r.font.name = "Calibri"
        sub_r.font.size = Pt(12)
        sub_r.font.italic = True
        sub_r.font.color.rgb = GRAY
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Meta table — bilingual labels
    now = datetime.now().strftime("%d/%m/%Y %H:%M")
    table_data = [
        ("Date / Ngày tạo",              now),
        ("Word count / Số từ",           f"{quality.get('word_count', 0):,}"),
        ("Sources / Số nguồn",           str(quality.get("source_count", 0))),
        ("Generation time / Thời gian",  f"{quality.get('total_latency_ms', 0)/1000:.1f}s"),
        ("Language / Ngôn ngữ",          "Bilingual — Vietnamese / English"),
        ("Report ID",                    rid[:20] + "..."),
    ]

    table = doc.add_table(rows=len(table_data), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)

    for row_idx, (label, value) in enumerate(table_data):
        row = table.rows[row_idx]
        label_cell = row.cells[0]
        value_cell = row.cells[1]

        label_cell.text = label
        label_cell.paragraphs[0].runs[0].font.bold = True
        label_cell.paragraphs[0].runs[0].font.color.rgb = NAVY
        set_cell_background(label_cell, "F1F5F9")

        value_cell.text = value
        value_cell.paragraphs[0].runs[0].font.color.rgb = GRAY

    # Page break after cover
    doc.add_page_break()


# ─── Main DOCX formatter ──────────────────────────────────────────────────────

class DocxFormatter:
    def generate(self, pipeline_result: dict, output_path: str) -> str:
        doc = Document()

        # Page setup: A4
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width  = Cm(21.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)

        configure_styles(doc)

        # ── Cover ────────────────────────────────────────────────────────────
        build_cover_page(doc, pipeline_result)

        # ── Main content ─────────────────────────────────────────────────────
        content = pipeline_result.get("markdown_content", "")
        parse_markdown_to_docx(doc, content)

        # ── Sources ───────────────────────────────────────────────────────────
        sources = pipeline_result.get("sources", [])
        if sources:
            doc.add_page_break()
            doc.add_heading("Tài liệu tham khảo (References)", level=1)
            for i, src in enumerate(sources, 1):
                title = src.get("title", "Không có tiêu đề")
                url   = src.get("url", "")
                snip  = src.get("snippet", "")[:150]
                score = src.get("relevance_score", 0)

                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                r = p.add_run(f"[{i}] {title}")
                r.bold = True

                if url:
                    p.add_run(f"\n{url[:80]}")
                    p.runs[-1].font.color.rgb = BLUE
                    p.runs[-1].font.italic = True

                if snip:
                    p.add_run(f"\n{snip}...")
                    p.runs[-1].font.color.rgb = GRAY

                p.add_run(f"  (relevance / độ liên quan: {score:.2f})")
                p.runs[-1].font.color.rgb = GRAY
                p.runs[-1].font.size = Pt(9)

        # ── Quality appendix ─────────────────────────────────────────────────
        quality = pipeline_result.get("quality", {})
        if quality:
            add_horizontal_rule(doc)
            h = doc.add_heading("Thông tin kỹ thuật (Technical Metadata)", level=2)
            h.runs[0].font.color.rgb = GRAY

            q_rows = [(k.replace("_", " ").title(), str(v))
                      for k, v in quality.items() if k != "step_timings_ms"]
            table = doc.add_table(rows=1 + len(q_rows), cols=2)
            table.style = "Light Shading Accent 1"

            # Header row
            header = table.rows[0]
            header.cells[0].text = "Chỉ số"
            header.cells[1].text = "Giá trị"
            for cell in header.cells:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_background(cell, "1B2B4B")
                cell.paragraphs[0].runs[0].font.color.rgb = WHITE

            for i, (k, v) in enumerate(q_rows, 1):
                row = table.rows[i]
                row.cells[0].text = k
                row.cells[1].text = v
                if i % 2 == 0:
                    set_cell_background(row.cells[0], "F1F5F9")
                    set_cell_background(row.cells[1], "F1F5F9")

        doc.save(output_path)
        return output_path